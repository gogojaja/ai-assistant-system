#!/usr/bin/env python3

"""
模块名称：test_office_integration
功能描述：TODO: 请补充功能描述
对外接口：
    - 无
依赖：
    - 标准库：os, pathlib, sys, tempfile
    - 第三方：core, docx
    - 项目内：无
版本：v1.0
更新记录：
    - 2026-05-23: 自动添加统一注释头
"""
"""集成测试：模拟飞书文件消息，验证文档处理链路"""

import sys
import os
import tempfile
from pathlib import Path

# 设置路径
sys.path.insert(0, str(Path.home() / "ai-assistant-system"))
sys.path.insert(0, str(Path.home() / "ai-assistant-system/assistants/office-assistant/src"))

from core.word_processor import WordProcessor
from core.summarizer import DocumentSummarizer
from docx import Document

# 创建测试文档
doc = Document()
doc.add_heading('集成测试报告', level=1)
doc.add_paragraph('本次测试验证 2号AI 的文档处理功能是否正常。')
doc.add_heading('测试环境', level=2)
doc.add_paragraph('Python 3.12.13, python-docx, Ollama 运行于本地。')
doc.add_heading('结论', level=2)
doc.add_paragraph('所有测试用例通过，系统运行稳定。')
test_path = tempfile.mktemp(suffix='.docx')
doc.save(test_path)

print("="*60)
print("开始集成测试...")

# 测试 WordProcessor
wp = WordProcessor(test_path)
info = wp.get_summary_info()
print(f"文档信息: {info}")

# 测试摘要
summarizer = DocumentSummarizer()
text = wp.extract_text()
result = summarizer.summarize(text, max_points=3)
print(f"\n摘要方法: {result['method']}")
print(f"摘要内容:\n{result['summary']}")

os.unlink(test_path)
print("\n✅ 集成测试完成，2号AI 核心功能正常")
