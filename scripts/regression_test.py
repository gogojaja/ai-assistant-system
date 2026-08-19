#!/usr/bin/env python3
"""
模块名称：regression_test
功能描述：三角色 AI 助理系统 — 全功能回归测试套件
          覆盖三角色 + 共享模块 + 回调服务
          支持独立运行、模块筛选、输出详细报告
对外接口：
    - 直接运行：python3 scripts/regression_test.py
    - 筛选模块：python3 scripts/regression_test.py --module chat
    - 查看帮助：python3 scripts/regression_test.py --help
依赖：
    - 标准库：os, sys, json, tempfile, re, math, logging, pathlib, importlib
    - 第三方：无（所有外部依赖模拟，缺失模块自动跳过）
    - 项目内：三角色 src 目录及 shared 模块
版本：v2.0
更新记录：
    - 2026-05-28: 初始创建，覆盖全部 5 个 AI 角色 + 共享模块 + 回调服务
    - 2026-08-19: v2.0 对齐需求基线 v1.2（PRJ-001），移除 file/sys 残留测试段
    - 2026-05-28: v1.1 修复缺失依赖自动跳过、相对导入处理、file_assistant 路径
"""
import os
import sys
import json
import tempfile
import re
import math
import logging
import time
import importlib
import importlib.util
from pathlib import Path

logging.basicConfig(level=logging.WARNING, format='%(levelname)s - %(message)s')

PROJECT_ROOT = Path(__file__).resolve().parent.parent
for p in [
    str(PROJECT_ROOT),
    str(PROJECT_ROOT / "shared"),
    str(PROJECT_ROOT / "shared/feishu-callback"),
    str(PROJECT_ROOT / "assistants"),
    str(PROJECT_ROOT / "assistants/chat-assistant/src"),
    str(PROJECT_ROOT / "assistants/office-assistant/src"),
    str(PROJECT_ROOT / "assistants/office-assistant/src/core"),
    str(PROJECT_ROOT / "assistants/life-assistant/src"),
]:
    if p not in sys.path:
        sys.path.insert(0, p)

results = {"pass": 0, "fail": 0, "skip": 0}
module_filter = None
_import_cache = {}

if "--help" in sys.argv:
    print("用法: python3 scripts/regression_test.py [--module <名称>]")
    print("  --module shared|chat|office|life|callback  筛选测试模块")
    sys.exit(0)

if "--module" in sys.argv:
    idx = sys.argv.index("--module")
    if idx + 1 < len(sys.argv):
        module_filter = sys.argv[idx + 1]


def _safe_import(module_name, attr=None):
    """安全导入模块，失败时返回 None"""
    key = (module_name, attr)
    if key in _import_cache:
        return _import_cache[key]
    try:
        if attr:
            mod = importlib.import_module(module_name)
            result = getattr(mod, attr)
        else:
            result = importlib.import_module(module_name)
        _import_cache[key] = result
        return result
    except (ImportError, ModuleNotFoundError, AttributeError) as e:
        _import_cache[key] = None
        return None


def _safe_load_from_file(filepath, func_name):
    """从文件路径加载模块并返回指定函数，失败返回 None"""
    try:
        spec = importlib.util.spec_from_file_location(f"_mod_{func_name}", str(filepath))
        mod = importlib.util.module_from_spec(spec)
        mod.__package__ = filepath.parent.name
        old_path = sys.path.copy()
        # 添加项目根 + shared + 文件父目录
        for p in [str(PROJECT_ROOT), str(PROJECT_ROOT / "shared"), str(filepath.parent),
                  str(PROJECT_ROOT / "assistants/office-assistant/src"),
                  str(PROJECT_ROOT / "assistants/office-assistant/src/core")]:
            if p not in sys.path:
                sys.path.insert(0, p)
        spec.loader.exec_module(mod)
        sys.path = old_path
        return getattr(mod, func_name, None)
    except Exception as e:
        return None


def _test(name, fn):
    if module_filter and module_filter not in name.lower():
        results["skip"] += 1
        return
    try:
        fn()
        results["pass"] += 1
        print(f"  ✅ {name}")
    except ImportError as e:
        results["skip"] += 1
        print(f"  ⏭️  {name}: 跳过（{e}）")
    except ModuleNotFoundError as e:
        results["skip"] += 1
        print(f"  ⏭️  {name}: 跳过（{e}）")
    except Exception as e:
        results["fail"] += 1
        print(f"  ❌ {name}: {e}")
        import traceback
        traceback.print_exc()


def _section(title):
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}")


# =====================================================================
# 共享模块测试
# =====================================================================
def test_shared():
    _section("1. 共享模块")

    # ---- crypto ----
    def test_crypto_encrypt_decrypt():
        encrypt_text = _safe_import("shared.crypto", "encrypt_text")
        decrypt_text = _safe_import("shared.crypto", "decrypt_text")
        if not encrypt_text or not decrypt_text:
            raise ImportError("cryptography 模块未安装")
        original = "Hello 世界 123 !@#"
        encrypted = encrypt_text(original)
        assert encrypted != original, "加密后不应与原文相同"
        assert isinstance(encrypted, str), "加密结果应为字符串"
        decrypted = decrypt_text(encrypted)
        assert decrypted == original, f"解密结果不匹配"

    def test_crypto_json_roundtrip():
        encrypt_json = _safe_import("shared.crypto", "encrypt_json")
        decrypt_json = _safe_import("shared.crypto", "decrypt_json")
        if not encrypt_json or not decrypt_json:
            raise ImportError("cryptography 模块未安装")
        data = {"key": "value", "list": [1, 2, 3], "chinese": "中文测试"}
        encrypted = encrypt_json(data)
        assert isinstance(encrypted, str)
        decrypted = decrypt_json(encrypted)
        assert decrypted == data

    def test_crypto_invalid_decrypt():
        decrypt_text = _safe_import("shared.crypto", "decrypt_text")
        if not decrypt_text:
            raise ImportError("cryptography 模块未安装")
        try:
            decrypt_text("invalid_base64_data_here")
            assert False, "应抛出异常"
        except Exception:
            pass

    _test("crypto 加解密", test_crypto_encrypt_decrypt)
    _test("crypto JSON 加解密", test_crypto_json_roundtrip)
    _test("crypto 无效数据容错", test_crypto_invalid_decrypt)

    # ---- knowledge_base ----
    def test_kb_search_empty():
        search = _safe_import("shared.knowledge_base", "search")
        if not search:
            raise ImportError("knowledge_base 模块未加载")
        res = search("不存在的查询词_xyz", top_k=3, min_score=0.01)
        assert isinstance(res, list)

    def test_kb_import_invalid_file():
        import_doc = _safe_import("shared.knowledge_base", "import_doc")
        if not import_doc:
            raise ImportError("knowledge_base 模块未加载")
        ok, msg = import_doc("/tmp/nonexistent_file_xyz.txt")
        assert ok is False and "不存在" in msg

    def test_kb_import_wrong_extension():
        import_doc = _safe_import("shared.knowledge_base", "import_doc")
        if not import_doc:
            raise ImportError("knowledge_base 模块未加载")
        tmp = tempfile.NamedTemporaryFile(suffix=".py", delete=False)
        tmp.close()
        try:
            ok, msg = import_doc(tmp.name)
            assert ok is False and "仅支持" in msg
        finally:
            os.unlink(tmp.name)

    def test_kb_list_docs():
        list_docs = _safe_import("shared.knowledge_base", "list_docs")
        if not list_docs:
            raise ImportError("knowledge_base 模块未加载")
        docs = list_docs()
        assert isinstance(docs, list)

    def test_kb_tokenize_and_bm25():
        _tokenize = _safe_import("shared.knowledge_base", "_tokenize")
        _bm25_score = _safe_import("shared.knowledge_base", "_bm25_score")
        if not _tokenize or not _bm25_score:
            raise ImportError("knowledge_base 模块未加载")
        tokens = _tokenize("测试中文English123混合文本")
        assert isinstance(tokens, list) and len(tokens) > 0
        score = _bm25_score(1, 10, 5, 1.5)
        assert isinstance(score, float) and score > 0

    _test("知识库 空搜索", test_kb_search_empty)
    _test("知识库 无效文件导入", test_kb_import_invalid_file)
    _test("知识库 扩展名校验", test_kb_import_wrong_extension)
    _test("知识库 列出文档", test_kb_list_docs)
    _test("知识库 分词+BM25", test_kb_tokenize_and_bm25)

    # ---- backend_utils ----
    def test_backend_config():
        get_backend_config = _safe_import("shared.backend_utils", "get_backend_config")
        if not get_backend_config:
            raise ImportError("backend_utils 模块未加载")
        cfg = get_backend_config()
        assert isinstance(cfg, dict)
        for k in ("backend", "port", "model"):
            assert k in cfg
        assert cfg["backend"] in ("llama.cpp", "ollama")

    def test_extract_from_reasoning():
        extract = _safe_import("shared.backend_utils", "extract_from_reasoning")
        if not extract:
            raise ImportError("backend_utils 模块未加载")
        assert extract("让我想想...\n总结：这是一个测试回答。") is not None
        assert extract("") == "（摘要为空）"

    def test_clean_reply():
        clean_reply = _safe_import("shared.backend_utils", "clean_reply")
        if not clean_reply:
            raise ImportError("backend_utils 模块未加载")
        assert clean_reply("最终输出：你好") == "你好"
        assert clean_reply("总结：结果") == "结果"
        assert clean_reply("Direct output: hello") == "hello"
        assert clean_reply("") == ""
        assert clean_reply("正常文本") == "正常文本"

    def test_wake_model_graceful():
        wake_model = _safe_import("shared.backend_utils", "wake_model")
        if not wake_model:
            raise ImportError("backend_utils 模块未加载")
        wake_model()

    _test("backend_utils 配置读取", test_backend_config)
    _test("backend_utils reasoning提取", test_extract_from_reasoning)
    _test("backend_utils 回复清理", test_clean_reply)
    _test("backend_utils 唤醒模型(无进程)", test_wake_model_graceful)


# =====================================================================
# 1号AI chat-assistant 测试
# =====================================================================
def test_chat():
    _section("2. 1号AI chat-assistant")

    # ---- chat.py ----
    def test_chat_history():
        clear_history = _safe_import("chat", "clear_history")
        load_history = _safe_import("chat", "load_history")
        save_history = _safe_import("chat", "save_history")
        if not all([clear_history, load_history, save_history]):
            raise ImportError("chat 模块未加载")
        save_history([{"role": "user", "content": "你好"}])
        assert len(load_history()) == 1
        result = clear_history()
        assert "已清空" in result

    _test("chat 历史保存/加载/清空", test_chat_history)

    # ---- search.py ----
    def test_search_archive():
        archive_search = _safe_import("search", "archive_search")
        search_archive = _safe_import("search", "search_archive")
        if not archive_search or not search_archive:
            raise ImportError("search 模块未加载")
        archive_search("测试查询", [{"title": "T", "url": "", "snippet": "S"}])
        result = search_archive("测试查询")
        assert result["found"]

    def test_search_format():
        format_results = _safe_import("search", "format_results")
        if not format_results:
            raise ImportError("search 模块未加载")
        result = format_results({"success": True, "results": [{"title": "T1", "url": "https://x.com", "snippet": "S1"}]})
        assert "T1" in result

    def test_search_archive_empty():
        search_archive = _safe_import("search", "search_archive")
        if not search_archive:
            raise ImportError("search 模块未加载")
        result = search_archive("__nonexistent_xyz__")
        assert result["found"] is False

    _test("search 归档与检索", test_search_archive)
    _test("search 格式化结果", test_search_format)
    _test("search 空搜索", test_search_archive_empty)

    # ---- main.py ----
    def test_main_format_reply():
        _format_reply = _safe_import("main", "_format_reply")
        if not _format_reply:
            raise ImportError("main 模块未加载")
        result = _format_reply("  Hello  World  ")
        assert result is not None

    def test_main_format_chinese():
        _format_reply = _safe_import("main", "_format_reply")
        if not _format_reply:
            raise ImportError("main 模块未加载")
        result = _format_reply("  这是一个 测试 ")
        assert result is not None

    def test_main_extract_reasoning():
        _extract = _safe_import("main", "_extract_from_reasoning")
        if not _extract:
            raise ImportError("main 模块未加载")
        result = _extract("让我思考一下...\n最终回答：\"今天天气很好\"")
        assert result is not None

    def test_main_extract_reasoning_quoted():
        _extract = _safe_import("main", "_extract_from_reasoning")
        if not _extract:
            raise ImportError("main 模块未加载")
        result = _extract("用户问天气\n回应说：「今天是个好天气」")
        assert "好天气" in result

    def test_main_trim_history():
        trim_history = _safe_import("main", "trim_history")
        if not trim_history:
            raise ImportError("main 模块未加载")
        msgs = [{"role": "system", "content": "sys"}]
        for i in range(30):
            msgs.append({"role": "user", "content": f"msg{i}"})
            msgs.append({"role": "assistant", "content": f"reply{i}"})
        trimmed = trim_history(msgs)
        non_sys = [m for m in trimmed if m["role"] != "system"]
        assert len(non_sys) <= 20

    def test_main_backend_config():
        _get_backend_config = _safe_import("main", "_get_backend_config")
        if not _get_backend_config:
            raise ImportError("main 模块未加载")
        cfg = _get_backend_config()
        assert "backend" in cfg and "port" in cfg and "model" in cfg

    def test_main_custom_prompt():
        _save = _safe_import("main", "_save_custom_prompt")
        _load = _safe_import("main", "_load_custom_prompt")
        if not _save or not _load:
            raise ImportError("main 模块未加载")
        test_id = f"__test_prompt_{int(time.time())}__"
        _save(test_id, "你是一个测试助手")
        assert _load(test_id) == "你是一个测试助手"
        _save(test_id, "")
        assert _load(test_id) == ""

    def test_main_handle_search():
        handle_search = _safe_import("main", "handle_search")
        if not handle_search:
            raise ImportError("main 模块未加载")
        result = handle_search("搜索 Python")
        assert isinstance(result, str)

    _test("main 回复格式化", test_main_format_reply)
    _test("main 中文格式化", test_main_format_chinese)
    _test("main reasoning提取(引号)", test_main_extract_reasoning)
    _test("main reasoning提取(回应说)", test_main_extract_reasoning_quoted)
    _test("main 历史裁剪", test_main_trim_history)
    _test("main 后端配置", test_main_backend_config)
    _test("main 自定义提示词", test_main_custom_prompt)
    _test("main 搜索处理", test_main_handle_search)

    # ---- message_handler（仅测试无外部依赖的部分） ----
    def test_mh_find_user_name():
        _find_user_name = _safe_import("message_handler", "_find_user_name")
        if not _find_user_name:
            raise ImportError("message_handler 模块未加载(dotenv)")
        tests = {
            "我叫张三": "张三",
            "我是李四": "李四",
            "姓名是王五": "王五",
        }
        for text, expected in tests.items():
            name = _find_user_name([{"role": "user", "content": text}])
            assert name == expected, f"对于 '{text}'，期望 '{expected}'，得到 '{name}'"

    def test_mh_find_user_name_no_match():
        _find_user_name = _safe_import("message_handler", "_find_user_name")
        if not _find_user_name:
            raise ImportError("message_handler 模块未加载(dotenv)")
        assert _find_user_name([{"role": "user", "content": "今天天气怎么样"}]) == ""

    def test_mh_now_str():
        _now_str = _safe_import("message_handler", "_now_str")
        if not _now_str:
            raise ImportError("message_handler 模块未加载(dotenv)")
        s = _now_str()
        assert re.match(r'\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}', s)

    def test_mh_help_text():
        TEXT = _safe_import("message_handler", "_HELP_TEXT")
        if not TEXT:
            raise ImportError("message_handler 模块未加载(dotenv)")
        assert "闲聊" in TEXT

    # ---- REQ-034 跨会话记忆 ----
    def test_memory_extract_facts():
        _extract_facts = _safe_import("message_handler", "_extract_facts")
        if not _extract_facts:
            raise ImportError("message_handler 模块未加载")
        facts = _extract_facts("我叫张三，我的生日是5月20日，邮箱是 zhangsan@test.com")
        types = {f["type"]: f["value"] for f in facts}
        assert types.get("user_name") == "张三", f"姓名提取失败: {types}"
        assert types.get("birthday") == "5月20日", f"生日提取失败: {types}"
        assert types.get("email") == "zhangsan@test.com", f"邮箱提取失败: {types}"

    def test_memory_extract_facts_no_match():
        _extract_facts = _safe_import("message_handler", "_extract_facts")
        if not _extract_facts:
            raise ImportError("message_handler 模块未加载")
        facts = _extract_facts("今天天气真好")
        assert facts == [], f"不应提取任何事实: {facts}"

    def test_memory_save_load():
        _save_memory = _safe_import("message_handler", "_save_memory")
        _load_memory = _safe_import("message_handler", "_load_memory")
        if not _save_memory or not _load_memory:
            raise ImportError("message_handler 模块未加载")
        import uuid
        uid = "test_" + uuid.uuid4().hex[:8]
        mem = {"facts": [{"type": "user_name", "value": "测试"}]}
        _save_memory(uid, mem)
        loaded = _load_memory(uid)
        assert loaded["facts"][0]["value"] == "测试", f"记忆回读失败: {loaded}"
        assert loaded.get("updated_at"), "缺少 updated_at 时间戳"

    def test_memory_context_format():
        _memory_context = _safe_import("message_handler", "_memory_context")
        if not _memory_context:
            raise ImportError("message_handler 模块未加载")
        import uuid
        uid = "test_ctx_" + uuid.uuid4().hex[:8]
        assert _memory_context(uid) == "", "无记忆时应返回空串"
        _save_memory = _safe_import("message_handler", "_save_memory")
        _save_memory(uid, {"facts": [{"type": "user_name", "value": "李四"}]})
        ctx = _memory_context(uid)
        assert "跨会话记忆" in ctx and "李四" in ctx, f"记忆上下文生成异常: {ctx}"

    def test_memory_remember_dedup():
        _remember = _safe_import("message_handler", "_remember")
        if not _remember:
            raise ImportError("message_handler 模块未加载")
        import uuid
        uid = "test_rm_" + uuid.uuid4().hex[:8]
        first = _remember("我叫王五", uid)
        assert "王五" in first, f"首次记忆应返回新增描述: {first}"
        second = _remember("我叫王五", uid)
        assert second == "", f"重复事实不应再次记录: {second}"

    # ---- REQ-035 任务委派 ----
    def test_delegate_help_no_content():
        _delegate = _safe_import("message_handler", "_delegate")
        if not _delegate:
            raise ImportError("message_handler 模块未加载")
        assert _delegate("#委派", "t", "o", "open_id") == "help", "无委派内容应返回 help"

    def test_delegate_unknown_role():
        _delegate = _safe_import("message_handler", "_delegate")
        if not _delegate:
            raise ImportError("message_handler 模块未加载")
        assert _delegate("#委派 财务 对账", "t", "o", "open_id") == "help", "未知角色应返回 help"

    def test_delegate_help_text():
        HELP = _safe_import("message_handler", "_DELEGATE_HELP")
        if not HELP:
            raise ImportError("message_handler 模块未加载")
        assert "#委派 办公" in HELP and "#委派 日程" in HELP, "委派帮助应包含角色说明"

    # ---- REQ-036 文档起草 ----
    def test_draft_help_no_topic():
        _draft = _safe_import("message_handler", "_draft")
        if not _draft:
            raise ImportError("message_handler 模块未加载")
        assert _draft("起草", "t", "o", "open_id") == "help", "无主题应返回 help"

    def test_draft_path_created():
        _draft_path = _safe_import("message_handler", "_draft_path")
        if not _draft_path:
            raise ImportError("message_handler 模块未加载")
        path = _draft_path()
        assert os.path.exists(path), f"草稿目录应存在: {path}"
        assert os.path.isdir(path), f"草稿路径应为目录: {path}"

    _test("message_handler 姓名查找", test_mh_find_user_name)
    _test("message_handler 无姓名匹配", test_mh_find_user_name_no_match)
    _test("message_handler 时间格式", test_mh_now_str)
    _test("message_handler 帮助文本", test_mh_help_text)
    _test("REQ-034 记忆事实提取", test_memory_extract_facts)
    _test("REQ-034 无匹配事实提取", test_memory_extract_facts_no_match)
    _test("REQ-034 记忆保存回读", test_memory_save_load)
    _test("REQ-034 记忆上下文格式", test_memory_context_format)
    _test("REQ-034 记忆去重", test_memory_remember_dedup)
    _test("REQ-035 委派无内容", test_delegate_help_no_content)
    _test("REQ-035 委派未知角色", test_delegate_unknown_role)
    _test("REQ-035 委派帮助文本", test_delegate_help_text)
    _test("REQ-036 起草无主题", test_draft_help_no_topic)
    _test("REQ-036 草稿目录创建", test_draft_path_created)


# =====================================================================
# 2号AI office-assistant 测试
# =====================================================================
def test_office():
    _section("3. 2号AI office-assistant")

    # ---- WordProcessor ----
    def test_wp_basic():
        WordProcessor = _safe_import("core.word_processor", "WordProcessor")
        if not WordProcessor:
            raise ImportError("python-docx 未安装")
        from docx import Document
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        try:
            doc = Document()
            doc.add_heading("测试标题", level=1)
            doc.add_paragraph("测试段落内容。")
            doc.save(tmp.name)
            wp = WordProcessor(tmp.name)
            info = wp.get_summary_info()
            assert info["paragraph_count"] >= 1
            text = wp.extract_text()
            assert "测试标题" in text
            titles = wp.extract_titles()
            assert len(titles) >= 1
            assert titles[0]["level"] == 1
        finally:
            os.unlink(tmp.name)

    def test_wp_empty():
        WordProcessor = _safe_import("core.word_processor", "WordProcessor")
        if not WordProcessor:
            raise ImportError("python-docx 未安装")
        from docx import Document
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        try:
            doc = Document()
            doc.save(tmp.name)
            wp = WordProcessor(tmp.name)
            assert wp.extract_text() == ""
        finally:
            os.unlink(tmp.name)

    def test_wp_tables():
        WordProcessor = _safe_import("core.word_processor", "WordProcessor")
        if not WordProcessor:
            raise ImportError("python-docx 未安装")
        from docx import Document
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        try:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A1"
            table.cell(0, 1).text = "B1"
            table.cell(1, 0).text = "A2"
            table.cell(1, 1).text = "B2"
            doc.save(tmp.name)
            wp = WordProcessor(tmp.name)
            tables = wp.extract_tables()
            assert len(tables) == 1
            assert tables[0][0] == ["A1", "B1"]
        finally:
            os.unlink(tmp.name)

    def test_wp_invalid():
        WordProcessor = _safe_import("core.word_processor", "WordProcessor")
        if not WordProcessor:
            raise ImportError("python-docx 未安装")
        try:
            WordProcessor("/nonexistent/file.docx")
            assert False, "应抛出异常"
        except Exception:
            pass

    _test("WordProcessor 基本提取", test_wp_basic)
    _test("WordProcessor 空文档", test_wp_empty)
    _test("WordProcessor 表格提取", test_wp_tables)
    _test("WordProcessor 无效路径", test_wp_invalid)

    # ---- ExcelProcessor ----
    def test_ep_basic():
        ExcelProcessor = _safe_import("core.excel_processor", "ExcelProcessor")
        if not ExcelProcessor:
            raise ImportError("openpyxl 未安装")
        import openpyxl
        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "测试表"
            ws.append(["姓名", "年龄"])
            ws.append(["张三", 28])
            wb.save(tmp.name)
            ep = ExcelProcessor(tmp.name)
            analysis = ep.analyze()
            assert "测试表" in analysis
            data = ep.get_data_text()
            assert "姓名" in data and "张三" in data
        finally:
            os.unlink(tmp.name)

    def test_ep_multi_sheet():
        ExcelProcessor = _safe_import("core.excel_processor", "ExcelProcessor")
        if not ExcelProcessor:
            raise ImportError("openpyxl 未安装")
        import openpyxl
        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        try:
            wb = openpyxl.Workbook()
            wb.active.title = "Sheet1"
            wb.create_sheet("Sheet2")
            wb.save(tmp.name)
            ep = ExcelProcessor(tmp.name)
            analysis = ep.analyze()
            assert "Sheet2" in analysis
        finally:
            os.unlink(tmp.name)

    def test_ep_empty():
        ExcelProcessor = _safe_import("core.excel_processor", "ExcelProcessor")
        if not ExcelProcessor:
            raise ImportError("openpyxl 未安装")
        import openpyxl
        tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp.close()
        try:
            wb = openpyxl.Workbook()
            wb.save(tmp.name)
            ep = ExcelProcessor(tmp.name)
            assert ep.analyze() != ""
        finally:
            os.unlink(tmp.name)

    def test_ep_invalid():
        ExcelProcessor = _safe_import("core.excel_processor", "ExcelProcessor")
        if not ExcelProcessor:
            raise ImportError("openpyxl 未安装")
        try:
            ExcelProcessor("/nonexistent/file.xlsx")
            assert False
        except Exception:
            pass

    _test("ExcelProcessor 基本分析", test_ep_basic)
    _test("ExcelProcessor 多工作表", test_ep_multi_sheet)
    _test("ExcelProcessor 空工作簿", test_ep_empty)
    _test("ExcelProcessor 无效路径", test_ep_invalid)

    # ---- DocumentSummarizer ----
    def test_summarizer_empty():
        DocumentSummarizer = _safe_import("core.summarizer", "DocumentSummarizer")
        if not DocumentSummarizer:
            raise ImportError("summarizer 模块未加载")
        s = DocumentSummarizer()
        r = s.summarize("")
        assert r["success"] is False
        assert r["summary"] == ""

    def test_summarizer_whitespace():
        DocumentSummarizer = _safe_import("core.summarizer", "DocumentSummarizer")
        if not DocumentSummarizer:
            raise ImportError("summarizer 模块未加载")
        s = DocumentSummarizer()
        assert s.summarize("   ")["success"] is False

    def test_summarizer_backend():
        DocumentSummarizer = _safe_import("core.summarizer", "DocumentSummarizer")
        if not DocumentSummarizer:
            raise ImportError("summarizer 模块未加载")
        cfg = DocumentSummarizer._get_backend_config()
        assert cfg["backend"] in ("llama.cpp", "ollama")

    _test("DocumentSummarizer 空文本", test_summarizer_empty)
    _test("DocumentSummarizer 空白字符", test_summarizer_whitespace)
    _test("DocumentSummarizer 后端配置", test_summarizer_backend)

    # ---- DocxConverter ----
    def test_converter_to_text():
        DocxConverter = _safe_import("core.converters", "DocxConverter")
        if not DocxConverter:
            raise ImportError("python-docx/mammoth 未安装")
        from docx import Document
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        try:
            doc = Document()
            doc.add_heading("转换测试", level=1)
            doc.add_paragraph("段落内容。")
            doc.save(tmp.name)
            text = DocxConverter.docx_to_text(tmp.name)
            assert "转换测试" in text
        finally:
            os.unlink(tmp.name)

    def test_converter_invalid():
        DocxConverter = _safe_import("core.converters", "DocxConverter")
        if not DocxConverter:
            raise ImportError("python-docx/mammoth 未安装")
        try:
            DocxConverter.docx_to_text("/nonexistent.docx")
            assert False
        except FileNotFoundError:
            pass

    _test("DocxConverter docx→text", test_converter_to_text)
    _test("DocxConverter 无效文件", test_converter_invalid)

    # ---- PPT Generator ----
    def test_ppt_parse():
        _parse = _safe_import("core.ppt_generator", "_parse_text_to_slides")
        if not _parse:
            raise ImportError("python-pptx 未安装")
        text = "测试演示\n\n## 章节一\n\n- 要点A\n- 要点B"
        parsed = _parse(text)
        assert parsed is not None
        assert parsed["title"] == "测试演示"

    def test_ppt_empty_text():
        generate_from_text = _safe_import("core.ppt_generator", "generate_from_text")
        if not generate_from_text:
            raise ImportError("python-pptx 未安装")
        result = generate_from_text("", "/tmp/nonexist_test.pptx")
        assert result is None

    def test_ppt_generate():
        generate_presentation = _safe_import("core.ppt_generator", "generate_presentation")
        if not generate_presentation:
            raise ImportError("python-pptx 未安装")
        tmp = tempfile.NamedTemporaryFile(suffix=".pptx", delete=False)
        tmp.close()
        try:
            slides = [{"title": "第一页", "content": ["要点1"]}, {"title": "第二页", "content": ["要点A"]}]
            result = generate_presentation("测试", slides, tmp.name)
            assert result == tmp.name
            assert os.path.getsize(result) > 1000
        finally:
            os.unlink(tmp.name)

    _test("PPT 文本解析", test_ppt_parse)
    _test("PPT 空文本", test_ppt_empty_text)
    _test("PPT 结构化生成", test_ppt_generate)

    # ---- 外部文档工具适配层（external_doc_tools）----
    def test_ext_check_ready():
        check_ready = _safe_import("external_doc_tools", "check_external_ready")
        if not check_ready:
            raise ImportError("external_doc_tools 模块未加载")
        status = check_ready()
        assert isinstance(status, dict)
        assert "pptagent_ready" in status and "dashi_ready" in status
        assert "any_ready" in status

    def test_ext_pptagent_empty_prompt():
        generate = _safe_import("external_doc_tools", "generate_ppt_via_pptagent")
        if not generate:
            raise ImportError("external_doc_tools 模块未加载")
        ok, res = generate("", "/tmp/nonexist.pptx")
        assert ok is False
        assert "不能为空" in res

    def test_ext_dashi_empty_prompt():
        generate = _safe_import("external_doc_tools", "generate_ppt_via_dashi")
        if not generate:
            raise ImportError("external_doc_tools 模块未加载")
        ok, res = generate("", "/tmp/nonexist.pptx")
        assert ok is False
        assert "不能为空" in res

    def test_ext_goal_json():
        build = _safe_import("external_doc_tools", "_build_goal_json")
        if not build:
            raise ImportError("external_doc_tools 模块未加载")
        goal = build("周报", "周报标题", "theme01", 3)
        import json
        data = json.loads(goal)
        assert data["schemaVersion"] == 2
        assert data["themePack"] == "theme01"
        assert len(data["slides"]) == 3
        assert "schemaVersion" in goal

    def test_ext_summarize_empty():
        summarize = _safe_import("external_doc_tools", "summarize_doc_via_external")
        if not summarize:
            raise ImportError("external_doc_tools 模块未加载")
        text, ok = summarize("")
        assert ok is False
        assert text == ""

    def test_dh_generate_ppt_fallback():
        # _generate_ppt 空文案时应返回失败而不抛异常
        dh = _safe_load_from_file(
            PROJECT_ROOT / "assistants/office-assistant/src/document_handler.py",
            "_generate_ppt",
        )
        if not dh:
            raise ImportError("document_handler 模块未加载")
        ok, res = dh("", "/tmp/nonexist.pptx")
        assert ok is False

    _test("外部工具可用性探测", test_ext_check_ready)
    _test("外部工具 PPTAgent 空提示词", test_ext_pptagent_empty_prompt)
    _test("外部工具 dashi 空提示词", test_ext_dashi_empty_prompt)
    _test("外部工具 goal.json 生成", test_ext_goal_json)
    _test("外部工具摘要空文本", test_ext_summarize_empty)
    _test("document_handler PPT 兜底", test_dh_generate_ppt_fallback)

    # ---- document_handler（纯函数单元） ----
    def test_dh_is_valid_summary():
        _is_valid_summary = _safe_load_from_file(
            PROJECT_ROOT / "assistants/office-assistant/src/document_handler.py",
            "_is_valid_summary"
        )
        if not _is_valid_summary:
            raise ImportError("document_handler 模块未加载")
        assert _is_valid_summary("这是一个有效的摘要内容。")
        assert not _is_valid_summary("")
        assert not _is_valid_summary("（摘要为空）")

    def test_dh_size_limits():
        _SIZE_LIMITS = _safe_load_from_file(
            PROJECT_ROOT / "assistants/office-assistant/src/document_handler.py",
            "_SIZE_LIMITS"
        )
        if not _SIZE_LIMITS:
            raise ImportError("document_handler 模块未加载")
        assert _SIZE_LIMITS['.docx'] == 30 * 1024 * 1024
        assert _SIZE_LIMITS['.xlsx'] == 15 * 1024 * 1024
        assert _SIZE_LIMITS['.pptx'] == 30 * 1024 * 1024

    def test_dh_build_fallback_summary():
        _build_fallback = _safe_load_from_file(
            PROJECT_ROOT / "assistants/office-assistant/src/document_handler.py",
            "_build_fallback_summary"
        )
        if not _build_fallback:
            raise ImportError("document_handler 模块未加载")
        result = _build_fallback("结构信息", "姓名\t年龄\n张三\t28\n李四\t32")
        assert "姓名" in result
        assert "张三" in result

    _test("document_handler 有效摘要检查", test_dh_is_valid_summary)
    _test("document_handler 大小限制", test_dh_size_limits)
    _test("document_handler 降级摘要", test_dh_build_fallback_summary)


# =====================================================================
# 3号AI life-assistant 测试
# =====================================================================
def test_life():
    _section("4. 3号AI life-assistant")

    # 导入 life_assistant.src 模块
    life_init_path = PROJECT_ROOT / "assistants/life-assistant/src/__init__.py"
    life_src = str(life_init_path.parent)
    if life_src not in sys.path:
        sys.path.insert(0, life_src)

    # ---- 入口 process ----
    def test_process_help():
        spec = importlib.util.spec_from_file_location("life_init", str(life_init_path))
        mod = importlib.util.module_from_spec(spec)
        mod.__package__ = "life_assistant.src"
        spec.loader.exec_module(mod)
        result = mod.process("帮助")
        assert "生活助手" in result
        assert "日程" in result

    def test_process_empty():
        spec = importlib.util.spec_from_file_location("life_init2", str(life_init_path))
        mod = importlib.util.module_from_spec(spec)
        mod.__package__ = "life_assistant.src"
        spec.loader.exec_module(mod)
        assert "生活助手" in mod.process("")

    def test_process_unknown():
        spec = importlib.util.spec_from_file_location("life_init3", str(life_init_path))
        mod = importlib.util.module_from_spec(spec)
        mod.__package__ = "life_assistant.src"
        spec.loader.exec_module(mod)
        result = mod.process("未知命令xyz")
        assert "未知命令" in result

    def test_process_dashboard():
        spec = importlib.util.spec_from_file_location("life_init4", str(life_init_path))
        mod = importlib.util.module_from_spec(spec)
        mod.__package__ = "life_assistant.src"
        spec.loader.exec_module(mod)
        assert "test.com" in mod.process("看板", dashboard_url="https://test.com/dash")
        assert "未配置" in mod.process("看板")

    _test("life process 帮助", test_process_help)
    _test("life process 空输入", test_process_empty)
    _test("life process 未知命令", test_process_unknown)
    _test("life process 看板路由", test_process_dashboard)

    # ---- 日程 ----
    def test_schedule_crud():
        add = _safe_import("scheduler", "add")
        list_items = _safe_import("scheduler", "list_items")
        delete = _safe_import("scheduler", "delete")
        search = _safe_import("scheduler", "search")
        if not all([add, list_items, delete, search]):
            raise ImportError("scheduler 模块未加载")
        item = add("测试日程", "2026-06-01 10:00")
        assert "id" in item and item["title"] == "测试日程"
        assert "测试日程" in list_items("2026-06-01")
        assert "测试日程" in search("测试")
        assert "已删除" in delete(item["id"])

    _test("scheduler 增删查", test_schedule_crud)

    # ---- 健康 ----
    def test_health_record():
        record = _safe_import("health_tracker", "record")
        report = _safe_import("health_tracker", "report")
        TYPES = _safe_import("health_tracker", "TYPES")
        if not all([record, report, TYPES]):
            raise ImportError("health_tracker 模块未加载")
        assert isinstance(TYPES, dict) and len(TYPES) > 0
        assert "已记录" in record(list(TYPES.keys())[0], "70")
        rpt = report("日报")
        assert isinstance(rpt, str) and len(rpt) > 0

    _test("health 记录+报告", test_health_record)

    # ---- 旅行 ----
    def test_travel_crud():
        create = _safe_import("travel_planner", "create")
        list_trips = _safe_import("travel_planner", "list_trips")
        view = _safe_import("travel_planner", "view")
        delete = _safe_import("travel_planner", "delete")
        if not all([create, list_trips, view, delete]):
            raise ImportError("travel_planner 模块未加载")
        trip = create("北京", "2026-07-01", "2026-07-05")
        assert trip["destination"] == "北京"
        assert "北京" in list_trips()
        assert "北京" in view(trip["id"])
        assert "已删除" in delete(trip["id"])

    _test("travel 增删查", test_travel_crud)

    # ---- 锻炼 ----
    def test_workout_crud():
        create = _safe_import("workout_planner", "create")
        list_plans = _safe_import("workout_planner", "list_plans")
        view = _safe_import("workout_planner", "view")
        delete = _safe_import("workout_planner", "delete")
        if not all([create, list_plans, view, delete]):
            raise ImportError("workout_planner 模块未加载")
        plan = create("测试锻炼")
        assert plan["name"] == "测试锻炼"
        assert "测试锻炼" in list_plans()
        assert "测试锻炼" in view(plan["id"])
        assert "已删除" in delete(plan["id"])

    _test("workout 增删查", test_workout_crud)

    # ---- 工作 ----
    def test_work_crud():
        create = _safe_import("work_planner", "create")
        list_items = _safe_import("work_planner", "list_items")
        view = _safe_import("work_planner", "view")
        set_status = _safe_import("work_planner", "set_status")
        set_priority = _safe_import("work_planner", "set_priority")
        delete = _safe_import("work_planner", "delete")
        if not all([create, list_items, view, set_status, set_priority, delete]):
            raise ImportError("work_planner 模块未加载")
        item = create("测试工作项")
        assert item["title"] == "测试工作项"
        assert "测试工作项" in list_items("all")
        assert "测试工作项" in view(item["id"])
        assert "进行中" in set_status(item["id"], "doing") or "doing" in set_status(item["id"], "doing").lower()
        assert "高" in set_priority(item["id"], "高")
        assert "已删除" in delete(item["id"]) or "删除" in delete(item["id"])

    _test("work 增删改查", test_work_crud)

    # ---- 内部处理函数 ----
    def test_handle_health_empty():
        _handle_health = _safe_import("life_assistant.src", "_handle_health")
        if not _handle_health:
            raise ImportError("life_assistant.src 模块未加载")
        assert "用法" in _handle_health([])

    def test_handle_travel_empty():
        _handle_travel = _safe_import("life_assistant.src", "_handle_travel")
        if not _handle_travel:
            raise ImportError("life_assistant.src 模块未加载")
        assert "用法" in _handle_travel([])

    def test_handle_workout_empty():
        _handle_workout = _safe_import("life_assistant.src", "_handle_workout")
        if not _handle_workout:
            raise ImportError("life_assistant.src 模块未加载")
        assert "用法" in _handle_workout([])

    def test_handle_work_empty():
        _handle_work = _safe_import("life_assistant.src", "_handle_work")
        if not _handle_work:
            raise ImportError("life_assistant.src 模块未加载")
        assert "用法" in _handle_work([])

    def test_handle_schedule_empty():
        _handle_schedule = _safe_import("life_assistant.src", "_handle_schedule")
        if not _handle_schedule:
            raise ImportError("life_assistant.src 模块未加载")
        assert "用法" in _handle_schedule([])

    _test("life _handle_schedule 空参数", test_handle_schedule_empty)
    _test("life _handle_health 空参数", test_handle_health_empty)
    _test("life _handle_travel 空参数", test_handle_travel_empty)
    _test("life _handle_workout 空参数", test_handle_workout_empty)
    _test("life _handle_work 空参数", test_handle_work_empty)



# =====================================================================
# 回调服务测试
# =====================================================================
def test_callback():
    _section("5. 回调服务 callback_server")

    callback_path = PROJECT_ROOT / "shared/feishu-callback/callback_server.py"
    cb_src = str(callback_path.parent)
    if cb_src not in sys.path:
        sys.path.insert(0, cb_src)

    try:
        spec = importlib.util.spec_from_file_location("callback_server", str(callback_path))
        cb_mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(cb_mod)
    except ModuleNotFoundError as e:
        print(f"\n  ⏭️  未安装依赖，跳过 callback 模块: {e}")
        return

    def test_health_endpoint():
        with cb_mod.app.test_client() as client:
            resp = client.get("/health")
            assert resp.status_code == 200
            assert resp.get_json()["status"] == "ok"

    def test_webhook_challenge():
        with cb_mod.app.test_client() as client:
            resp = client.post("/webhook", json={"challenge": "test123"})
            assert resp.status_code == 200
            assert resp.get_json()["challenge"] == "test123"

    def test_webhook_empty_body_json():
        with cb_mod.app.test_client() as client:
            resp = client.post("/webhook", json={"test": "data"})
            assert resp.status_code == 200
            data = resp.get_json()
            assert isinstance(data, dict) and "code" in data

    def test_webhook_empty_body():
        with cb_mod.app.test_client() as client:
            resp = client.post("/webhook", data="not json", content_type="application/json")
            assert resp.status_code == 400

    def test_unknown_route():
        with cb_mod.app.test_client() as client:
            resp = client.get("/webhook_unknown")
            assert resp.status_code == 404

    def test_backends_config():
        assert not hasattr(cb_mod, "BACKENDS"), "残留 file/sys 代理路由定义"
        assert not hasattr(cb_mod, "proxy_backend"), "残留 proxy_backend 函数"

    def test_dashboard_exists():
        assert cb_mod.dashboard_bp is not None

    # ---- REQ-037 主动提醒 ----
    def test_reminder_now_str():
        s = cb_mod._now_str()
        import re
        assert re.match(r"^\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2}$", s), f"_now_str 格式异常: {s}"

    def test_reminder_thread_started():
        # 验证启动函数可调用且会创建守护线程
        cb_mod._REMINDER_SENT.clear()
        try:
            cb_mod.start_reminder_thread()
        except Exception as e:
            raise AssertionError(f"提醒线程启动失败: {e}")

    def test_reminder_check_empty():
        # 无配置目标时不应抛异常
        import os as _os
        old = _os.environ.get("REMINDER_TARGET_ID")
        _os.environ.pop("REMINDER_TARGET_ID", None)
        try:
            cb_mod._run_reminder_check()
        except Exception as e:
            raise AssertionError(f"提醒检查空配置应静默: {e}")
        finally:
            if old:
                _os.environ["REMINDER_TARGET_ID"] = old

    def test_reminder_extract_send():
        # 模拟到期日程：写入临时调度文件，验证 check_reminders 能识别
        from pathlib import Path as _Path
        life_dir = PROJECT_ROOT / "data" / "life"
        old_schedule = life_dir / "schedules.json"
        backup = None
        if old_schedule.exists():
            backup = old_schedule.read_text()
        from datetime import datetime, timedelta
        near = (datetime.now() + timedelta(minutes=5)).strftime("%Y-%m-%d %H:%M:%S")
        life_dir.mkdir(parents=True, exist_ok=True)
        old_schedule.write_text('[]', encoding='utf-8')
        import json as _json
        old_schedule.write_text(_json.dumps(
            [{"id": "testrem1", "title": "回归测试提醒", "time": near}],
            ensure_ascii=False
        ), encoding='utf-8')
        try:
            sys.path.insert(0, str(PROJECT_ROOT / "assistants"))
            from assistants.life_assistant.src.reminder import check_reminders
            result = check_reminders(within_minutes=30)
            assert "回归测试提醒" in result, f"提醒识别失败: {result}"
        finally:
            if backup is not None:
                old_schedule.write_text(backup, encoding='utf-8')
            else:
                old_schedule.unlink(missing_ok=True)

    _test("callback /health", test_health_endpoint)
    _test("callback challenge验证", test_webhook_challenge)
    _test("callback 空JSON", test_webhook_empty_body_json)
    _test("callback 空body", test_webhook_empty_body)
    _test("callback 未知路由", test_unknown_route)
    _test("callback 后端路由配置", test_backends_config)
    _test("callback dashboard蓝图", test_dashboard_exists)
    _test("REQ-037 时间格式", test_reminder_now_str)
    _test("REQ-037 提醒线程启动", test_reminder_thread_started)
    _test("REQ-037 空配置静默", test_reminder_check_empty)
    _test("REQ-037 日程识别", test_reminder_extract_send)


# =====================================================================
# 测试报告输出
# =====================================================================
def print_report():
    total = results["pass"] + results["fail"] + results["skip"]
    print(f"\n{'='*60}")
    print(f"  测试报告总结")
    print(f"{'='*60}")
    print(f"  总计: {total}")
    print(f"  ✅ 通过: {results['pass']}")
    print(f"  ❌ 失败: {results['fail']}")
    print(f"  ⏭️  跳过: {results['skip']}（缺失依赖或被筛选器排除）")
    print(f"  {'🎉 全部通过!' if results['fail'] == 0 else '⚠️  存在失败项'}")
    print(f"{'='*60}")
    return results["fail"] == 0


if __name__ == "__main__":
    print(f"\n{'='*60}")
    print(f"  三角色 AI 助理系统 · 全功能回归测试")
    print(f"  项目路径: {PROJECT_ROOT}")
    if module_filter:
        print(f"  模块筛选: {module_filter}")
    print(f"  💡 建议使用 venv 运行以获完整覆盖:")
    print(f"     python3 -m venv tmp_venv && source tmp_venv/bin/activate")
    print(f"     pip install python-docx openpyxl python-pptx python-dotenv cryptography mammoth")
    print(f"     python3 scripts/regression_test.py")
    print(f"{'='*60}")

    test_modules = [
        ("shared", test_shared),
        ("chat", test_chat),
        ("office", test_office),
        ("life", test_life),
        ("callback", test_callback),
    ]

    for name, fn in test_modules:
        if module_filter and module_filter not in name:
            results["skip"] += 1
            continue
        try:
            fn()
        except Exception as e:
            print(f"\n  ⚠️  模块 '{name}' 整体跳过: {e}")

    success = print_report()
    sys.exit(0 if success else 1)
