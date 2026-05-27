#!/usr/bin/env python3

"""
模块名称：word_processor
功能描述：Word 文档处理（提取文本、摘要信息、标题）
对外接口：
    - WordProcessor(file_path): 初始化并解析文档
        - get_summary_info(): 获取文档基本信息
        - extract_text(): 提取纯文本
        - extract_titles(): 提取标题列表
依赖：
    - 标准库：logging
    - 第三方：python-docx, mammoth
    - 项目内：无
版本：v1.0
更新记录：
    - 2026-05-23: 添加统一注释头
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger("WordProcessor")
logger.setLevel(logging.DEBUG)
if not logger.handlers:
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.DEBUG)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)

# 文件处理器（可选，后续写入文件）
# file_handler = logging.FileHandler('logs/office_assistant.log')
# file_handler.setLevel(logging.DEBUG)
# file_handler.setFormatter(formatter)
# logger.addHandler(file_handler)


class WordProcessor:
    """Word 文档处理类"""
    
    def __init__(self, file_path: str):
        """
        初始化处理器
        :param file_path: .docx 文件路径
        """
        self.file_path = Path(file_path)
        self.document: Optional[Document] = None
        self._load_document()

    def _load_document(self):
        """加载 Word 文档，带异常处理"""
        try:
            if not self.file_path.exists():
                raise FileNotFoundError(f"文件不存在: {self.file_path}")
            if self.file_path.suffix.lower() != '.docx':
                raise ValueError(f"不支持的文件格式: {self.file_path.suffix}，仅支持 .docx")
            
            self.document = Document(str(self.file_path))
            logger.info(f"成功加载文档: {self.file_path.name}")
        except PackageNotFoundError:
            logger.error("文件损坏或不是有效的 .docx 格式")
            raise
        except Exception as e:
            logger.error(f"加载文档失败: {e}")
            raise

    def extract_text(self) -> str:
        """
        提取文档全部纯文本（段落 + 表格内文本）
        :return: 合并后的文本字符串
        """
        if not self.document:
            raise RuntimeError("文档未加载")
        
        text_parts = []
        
        # 提取段落文本
        for para in self.document.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # 提取表格文本
        for table in self.document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        full_text = '\n'.join(text_parts)
        logger.debug(f"提取文本总长度: {len(full_text)} 字符")
        return full_text

    def extract_titles(self) -> List[Dict[str, str]]:
        """
        提取文档标题及层级（识别 Word 内置标题样式 Heading 1~6）
        :return: 标题列表，每个元素包含 'level' 和 'text'
        """
        if not self.document:
            raise RuntimeError("文档未加载")
        
        titles = []
        for para in self.document.paragraphs:
            style_name = para.style.name if para.style else ''
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1  # 默认为一级标题
                titles.append({
                    'level': level,
                    'text': para.text.strip()
                })
        
        logger.info(f"提取到 {len(titles)} 个标题")
        for title in titles:
            logger.debug(f"标题 Lv{title['level']}: {title['text']}")
        return titles

    def extract_tables(self) -> List[List[List[str]]]:
        """
        提取所有表格数据（二维列表）
        :return: 表格列表，每个表格为行列表，每行为单元格字符串列表
        """
        if not self.document:
            raise RuntimeError("文档未加载")
        
        tables_data = []
        for idx, table in enumerate(self.document.tables):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_cells)
            tables_data.append(table_rows)
            logger.debug(f"表格 {idx+1}: {len(table_rows)} 行 x {len(table_rows[0]) if table_rows else 0} 列")
        
        logger.info(f"提取到 {len(tables_data)} 个表格")
        return tables_data

    def get_summary_info(self) -> Dict:
        """
        获取文档概要信息
        :return: 包含文件名、段落数、标题数、表格数等信息的字典
        """
        if not self.document:
            raise RuntimeError("文档未加载")
        
        summary = {
            'file_name': self.file_path.name,
            'paragraph_count': len(self.document.paragraphs),
            'title_count': len(self.extract_titles()),
            'table_count': len(self.document.tables),
            'character_count': len(self.extract_text())
        }
        logger.info(f"文档概要: {summary}")
        return summary


# 命令行入口（用于直接测试）
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("使用方法: python word_processor.py <word文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    wp = WordProcessor(file_path)
    
    print("="*60)
    print("文档概要:")
    summary = wp.get_summary_info()
    for k, v in summary.items():
        print(f"  {k}: {v}")
    
    print("\n" + "="*60)
    print("标题结构:")
    for title in wp.extract_titles():
        indent = "  " * (title['level'] - 1)
        print(f"{indent}- {title['text']}")
    
    print("\n" + "="*60)
    print("全部文本（前500字符）:")
    text = wp.extract_text()
    print(text[:500] + ("..." if len(text) > 500 else ""))
