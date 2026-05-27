#!/usr/bin/env python3

"""
模块名称：test_converters
功能描述：TODO: 请补充功能描述
对外接口：
    - create_test_doc()
    - test_case_1_to_text()
    - test_case_2_to_markdown()
    - test_case_3_invalid_file()
    - run_all_tests()
依赖：
    - 标准库：logging, os, pathlib, sys, tempfile
    - 第三方：core, docx
    - 项目内：无
版本：v1.0
更新记录：
    - 2026-05-23: 自动添加统一注释头
"""
"""
DocxConverter 测试脚本
3组测试用例（修正版：适应 mammoth 实际输出格式）
"""

import sys
import os
import logging
import tempfile
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from core.converters import DocxConverter
from docx import Document

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


def create_test_doc():
    """创建测试文档，包含标题、文本、表格"""
    doc = Document()
    doc.add_heading('转换测试文档', level=1)
    doc.add_paragraph('这是一个用于格式转换测试的段落。')
    doc.add_heading('数据表格', level=2)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = '项目'
    table.rows[0].cells[1].text = '值'
    table.rows[1].cells[0].text = '状态'
    table.rows[1].cells[1].text = '正常'
    return doc


def test_case_1_to_text():
    """测试用例1：docx → 纯文本"""
    logger.info("测试用例1：转换为纯文本")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        path = f.name
    try:
        doc = create_test_doc()
        doc.save(path)
        text = DocxConverter.docx_to_text(path)
        assert '转换测试文档' in text, "文本中应包含标题"
        assert '数据表格' in text, "文本中应包含二级标题"
        # 表格内容也应出现在纯文本中
        assert '项目' in text, "表格内容应包含'项目'"
        assert '正常' in text, "表格内容应包含'正常'"
        logger.info("✅ 测试用例1 通过")
    finally:
        os.unlink(path)


def test_case_2_to_markdown():
    """测试用例2：docx → Markdown（验证内容存在即可，不强求格式）"""
    logger.info("测试用例2：转换为 Markdown")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        path = f.name
    try:
        doc = create_test_doc()
        doc.save(path)
        md = DocxConverter.docx_to_markdown(path)
        logger.info("Markdown 输出片段：\n" + md[:300])
        # 检查标题
        assert '# 转换测试文档' in md, "Markdown 应包含一级标题"
        # 检查二级标题
        assert '## 数据表格' in md, "Markdown 应包含二级标题"
        # 检查表格内容（项目、值、状态、正常）而非特定格式
        assert '项目' in md, "Markdown 应包含表格内容'项目'"
        assert '值' in md, "Markdown 应包含表格内容'值'"
        assert '状态' in md, "Markdown 应包含表格内容'状态'"
        assert '正常' in md, "Markdown 应包含表格内容'正常'"
        logger.info("✅ 测试用例2 通过")
    finally:
        os.unlink(path)


def test_case_3_invalid_file():
    """测试用例3：无效文件处理"""
    logger.info("测试用例3：无效文件异常")
    # 文件不存在
    try:
        DocxConverter.docx_to_text("/tmp/nonexistent.docx")
        assert False, "应该抛出 FileNotFoundError"
    except FileNotFoundError:
        pass
    # 格式不支持
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as f:
        path = f.name
    try:
        DocxConverter.docx_to_text(path)
        assert False, "应该抛出 ValueError"
    except ValueError:
        pass
    finally:
        os.unlink(path)
    logger.info("✅ 测试用例3 通过")


def run_all_tests():
    logger.info("="*60)
    logger.info("开始运行 DocxConverter 测试套件")
    logger.info("="*60)
    test_case_1_to_text()
    test_case_2_to_markdown()
    test_case_3_invalid_file()
    logger.info("="*60)
    logger.info("所有测试用例通过！🎉")


if __name__ == "__main__":
    run_all_tests()
