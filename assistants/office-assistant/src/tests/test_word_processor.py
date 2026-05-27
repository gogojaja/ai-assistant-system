#!/usr/bin/env python3

"""
模块名称：test_word_processor
功能描述：TODO: 请补充功能描述
对外接口：
    - create_test_docx()
    - test_case_1_basic_extraction()
    - test_case_2_empty_document()
    - test_case_3_file_not_found()
    - run_all_tests()
依赖：
    - 标准库：logging, os, pathlib, sys, tempfile
    - 第三方：core, docx
    - 项目内：无
版本：v1.0
更新记录：
    - 2026-05-23: 自动添加统一注释头
"""
"""
WordProcessor 单元测试脚本
内置 3 组标准测试用例，可一键运行
"""

import sys
import os
import logging
import tempfile
from pathlib import Path

# 将 src 目录加入路径，以便导入 core 模块
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from core.word_processor import WordProcessor
from docx import Document

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


def create_test_docx(file_path: str):
    """创建一个用于测试的 Word 文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('测试文档', level=1)
    doc.add_paragraph('这是一个用于测试的段落，包含一些文本内容。')
    
    doc.add_heading('第一部分：简介', level=2)
    doc.add_paragraph('这是简介段落，用于验证段落提取功能。')
    
    doc.add_heading('第二部分：详细数据', level=2)
    
    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '姓名'
    hdr_cells[1].text = '年龄'
    hdr_cells[2].text = '城市'
    row1_cells = table.rows[1].cells
    row1_cells[0].text = '张三'
    row1_cells[1].text = '28'
    row1_cells[2].text = '北京'
    row2_cells = table.rows[2].cells
    row2_cells[0].text = '李四'
    row2_cells[1].text = '32'
    row2_cells[2].text = '上海'
    
    doc.add_paragraph('')  # 空行
    doc.add_paragraph('文档结束。')
    
    doc.save(file_path)
    logger.info(f"测试文档已创建: {file_path}")


def test_case_1_basic_extraction():
    """
    测试用例1：基本文本、标题、表格提取
    验证：能正确提取文本、标题数量和表格数据
    """
    logger.info("开始测试用例1：基本提取功能")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = f.name
    
    try:
        create_test_docx(temp_path)
        wp = WordProcessor(temp_path)
        
        # 验证文本提取
        text = wp.extract_text()
        assert '测试文档' in text, "文本提取失败：未找到标题"
        assert '张三' in text, "文本提取失败：未找到表格内容"
        assert '北京' in text, "文本提取失败：未找到表格城市"
        
        # 验证标题提取
        titles = wp.extract_titles()
        assert len(titles) == 3, f"标题数量错误，期望3，实际{len(titles)}"
        assert titles[0]['text'] == '测试文档'
        assert titles[0]['level'] == 1
        assert titles[1]['text'] == '第一部分：简介'
        assert titles[1]['level'] == 2
        
        # 验证表格提取
        tables = wp.extract_tables()
        assert len(tables) == 1, f"表格数量错误，期望1，实际{len(tables)}"
        assert tables[0][0] == ['姓名', '年龄', '城市']  # 表头
        assert tables[0][1] == ['张三', '28', '北京']
        
        logger.info("✅ 测试用例1 通过")
    finally:
        os.unlink(temp_path)


def test_case_2_empty_document():
    """
    测试用例2：空文档处理
    验证：加载空白文档不应报错，提取内容为空
    """
    logger.info("开始测试用例2：空文档处理")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = f.name
    
    try:
        doc = Document()
        doc.save(temp_path)
        
        wp = WordProcessor(temp_path)
        text = wp.extract_text()
        assert text == '', f"空文档文本应为空，实际长度{len(text)}"
        
        titles = wp.extract_titles()
        assert len(titles) == 0, f"空文档标题应为0，实际{len(titles)}"
        
        tables = wp.extract_tables()
        assert len(tables) == 0, f"空文档表格应为0，实际{len(tables)}"
        
        logger.info("✅ 测试用例2 通过")
    finally:
        os.unlink(temp_path)


def test_case_3_file_not_found():
    """
    测试用例3：文件不存在或格式错误
    验证：抛出明确异常
    """
    logger.info("开始测试用例3：异常处理")
    
    # 文件不存在
    try:
        wp = WordProcessor("/non/existent/file.docx")
        assert False, "应该抛出 FileNotFoundError"
    except FileNotFoundError:
        logger.info("  - 文件不存在异常捕获成功")
    except Exception as e:
        assert False, f"预期 FileNotFoundError，实际 {type(e).__name__}"
    
    # 格式不支持（创建临时 .txt 文件模拟）
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as f:
        f.write(b"not a docx")
        temp_path = f.name
    try:
        wp = WordProcessor(temp_path)
        assert False, "应该抛出 ValueError"
    except ValueError:
        logger.info("  - 格式错误异常捕获成功")
    except Exception as e:
        assert False, f"预期 ValueError，实际 {type(e).__name__}"
    finally:
        os.unlink(temp_path)
    
    logger.info("✅ 测试用例3 通过")


def run_all_tests():
    """一键运行所有测试用例"""
    logger.info("="*60)
    logger.info("开始运行 WordProcessor 测试套件")
    logger.info("="*60)
    
    test_case_1_basic_extraction()
    test_case_2_empty_document()
    test_case_3_file_not_found()
    
    logger.info("="*60)
    logger.info("所有测试用例通过！🎉")
    logger.info("="*60)


if __name__ == "__main__":
    run_all_tests()
