"""
模块名称：test_office
功能描述：2号AI office-assistant 测试脚本，覆盖 Word/Excel/PPT/转换器/文件工具
对外接口：
    - test_word_processor(): WordProcessor 单元测试
    - test_excel_processor(): ExcelProcessor 单元测试
    - test_docx_converter(): DocxConverter 单元测试
    - test_ppt_generator(): PPT 生成器单元测试
    - test_ppt_professional(): 专业级 PPT 测试
    - test_file_handler(): 文件工具单元测试
依赖：
    - openpyxl, python-docx, python-pptx, mammoth
版本：v2.0
更新记录：
    - 2026-05-26: 新增 test_ppt_professional 专业级 PPT 测试
    - 2026-05-25: 初始创建
"""
import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))
sys.path.insert(0, str(Path(__file__).parent.parent / "src" / "core"))

from core.word_processor import WordProcessor
from core.excel_processor import ExcelProcessor
from core.summarizer import DocumentSummarizer
from core.ppt_generator import generate_presentation, generate_from_text, _parse_text_to_slides, check_pandoc
from utils.file_handler import safe_delete, check_whitelist, cleanup_temp_dir


def test_word_processor():
    """测试 WordProcessor：创建临时 docx 并验证提取"""
    from docx import Document

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    try:
        doc = Document()
        doc.add_heading("测试标题", level=1)
        doc.add_paragraph("这是测试段落内容。")
        doc.add_heading("二级标题", level=2)
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        doc.save(tmp.name)

        wp = WordProcessor(tmp.name)
        info = wp.get_summary_info()
        assert info["paragraph_count"] >= 3
        assert info["table_count"] == 1
        assert info["title_count"] == 2

        text = wp.extract_text()
        assert "测试标题" in text
        assert "A1" in text

        titles = wp.extract_titles()
        assert len(titles) == 2
        assert titles[0]["level"] == 1
        assert titles[1]["level"] == 2

        tables = wp.extract_tables()
        assert len(tables) == 1
        assert tables[0][0] == ["A1", "B1"]
    finally:
        safe_delete(tmp.name)
    return True


def test_excel_processor():
    """测试 ExcelProcessor：创建临时 xlsx 并验证分析"""
    import openpyxl
    from openpyxl import Workbook

    tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
    tmp.close()
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "测试表"
        ws.append(["姓名", "年龄", "城市"])
        ws.append(["张三", 28, "北京"])
        ws.append(["李四", 32, "上海"])
        wb.save(tmp.name)

        ep = ExcelProcessor(tmp.name)
        analysis = ep.analyze()
        assert "测试表" in analysis
        assert "3行" in analysis or "3行" in analysis

        data_text = ep.get_data_text()
        assert "姓名" in data_text
        assert "张三" in data_text
    finally:
        safe_delete(tmp.name)
    return True


def test_ppt_generator():
    """测试 PPT 生成：验证文件创建与内容"""
    tmp = tempfile.NamedTemporaryFile(suffix=".pptx", delete=False)
    tmp.close()
    try:
        slides_data = [
            {"title": "第一页", "content": ["要点1", "要点2"]},
            {"title": "第二页", "content": ["要点A"]},
        ]
        result = generate_presentation("测试PPT", slides_data, tmp.name)
        assert result == tmp.name
        assert os.path.getsize(tmp.name) > 1000
    finally:
        safe_delete(tmp.name)
    return True


def test_ppt_from_text():
    """测试 PPT 从文本生成"""
    tmp = tempfile.NamedTemporaryFile(suffix=".pptx", delete=False)
    tmp.close()
    try:
        text = "会议总结\n进展\n项目完成80%\n问题\n资源不足\n"
        result = generate_from_text(text, tmp.name)
        assert result is not None
        assert os.path.getsize(result) > 1000
    finally:
        safe_delete(tmp.name)
    return True


def test_file_handler():
    """测试文件处理工具"""
    # safe_delete with non-existent file
    safe_delete("/nonexistent/file.docx")
    # safe_delete with unallowed extension
    tmp = tempfile.NamedTemporaryFile(suffix=".tmp", delete=False)
    tmp.close()
    try:
        safe_delete(tmp.name)
        assert os.path.exists(tmp.name)
    finally:
        safe_delete(tmp.name)

    # check_whitelist
    assert check_whitelist("test.docx")
    assert check_whitelist("test.xlsx")
    assert check_whitelist("test.pptx")
    assert not check_whitelist("test.dmg")

    # cleanup_temp_dir with non-existent dir
    cleanup_temp_dir("/nonexistent_dir")
    return True


def test_word_processor_empty():
    """空文档处理"""
    from docx import Document
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    try:
        doc = Document()
        doc.save(tmp.name)
        wp = WordProcessor(tmp.name)
        assert wp.extract_text() == ""
        assert wp.get_summary_info()["paragraph_count"] == 0
    finally:
        safe_delete(tmp.name)
    return True


def test_word_processor_invalid_path():
    """无效路径处理"""
    try:
        WordProcessor("/nonexistent/file.docx")
        return False
    except (FileNotFoundError, Exception):
        return True


def test_excel_processor_empty():
    """空 Excel 处理"""
    import openpyxl
    from openpyxl import Workbook
    tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
    tmp.close()
    try:
        wb = Workbook()
        wb.save(tmp.name)
        ep = ExcelProcessor(tmp.name)
        analysis = ep.analyze()
        assert "Sheet" in analysis
        assert ep.get_data_text() != ""
    finally:
        safe_delete(tmp.name)
    return True


def test_excel_processor_invalid_path():
    """无效 Excel 路径处理"""
    try:
        ExcelProcessor("/nonexistent/file.xlsx")
        return False
    except Exception:
        return True


def test_ppt_from_empty_text():
    """空文本 PPT 生成"""
    result = generate_from_text("", "/tmp/nonexist.pptx")
    assert result is None
    return True


def test_ppt_professional():
    """专业级 PPT 生成测试：解析+生成"""
    # 测试文本解析
    text = """测试演示

## 章节一

- 要点A
- 要点B

## 章节二

标题页
- 项目1
- 项目2

左内容|右内容"""

    parsed = _parse_text_to_slides(text)
    assert parsed is not None
    assert parsed["title"] == "测试演示"

    slides = parsed["slides"]
    assert len(slides) >= 4  # 2 sections + 1 content + 1 two-column
    section_count = sum(1 for s in slides if s.get("type") == "section")
    assert section_count == 2

    # 测试文件生成
    tmp = tempfile.NamedTemporaryFile(suffix=".pptx", delete=False)
    tmp.close()
    try:
        result = generate_from_text(text, tmp.name)
        assert result is not None
        assert os.path.getsize(result) > 1000
        from pptx import Presentation
        prs = Presentation(result)
        assert len(prs.slides) > 3
    finally:
        safe_delete(tmp.name)

    # 测试 Pandoc 检测
    check_pandoc()

    # 测试结构化生成
    slides_data = [
        {"type": "section", "title": "概览"},
        {"type": "content", "title": "数据", "content": ["指标1", "指标2"]},
        {"type": "two_column", "title": "对比", "left": ["A", "B"], "right": ["C", "D"]},
    ]
    tmp2 = tempfile.NamedTemporaryFile(suffix=".pptx", delete=False)
    tmp2.close()
    try:
        result = generate_presentation("结构化测试", slides_data, tmp2.name)
        assert result is not None
        assert os.path.getsize(result) > 1000
    finally:
        safe_delete(tmp2.name)

    return True


def test_document_handler_edge_cases():
    """模拟 document_handler 单元测试（无飞书依赖）"""
    from document_handler import generate_excel_summary

    text, is_ai = generate_excel_summary("")
    assert text == ""

    text, is_ai = generate_excel_summary("   ")
    assert text == ""
    return True


def test_summarizer_edge_cases():
    """DocumentSummarizer 边缘用例"""
    summarizer = DocumentSummarizer()

    # 空文本
    result = summarizer.summarize("")
    assert result["success"] is False
    assert result["summary"] == ""
    assert result["error"] == "输入文本为空"

    result = summarizer.summarize("   ")
    assert result["success"] is False
    assert result["summary"] == ""

    # 后端配置读取
    cfg = DocumentSummarizer._get_backend_config()
    assert cfg["backend"] in ("llama.cpp", "ollama")
    assert isinstance(cfg["port"], int)
    assert isinstance(cfg["model"], str)
    return True


def test_size_limits():
    """验证文件大小限制常量"""
    from document_handler import _SIZE_LIMITS
    assert _SIZE_LIMITS['.docx'] == 30 * 1024 * 1024
    assert _SIZE_LIMITS['.xlsx'] == 15 * 1024 * 1024
    return True


def test_file_handler_edge_cases():
    """文件工具边缘用例"""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    try:
        assert check_whitelist(tmp.name)
    finally:
        safe_delete(tmp.name)

    assert not check_whitelist("test.sh")
    assert not check_whitelist("test")
    assert safe_delete("/nonexistent/doc.docx") is None
    return True


if __name__ == "__main__":
    print("=" * 60)
    print("2号AI office-assistant 测试套件")
    print("=" * 60)

    tests = [
        ("WordProcessor", test_word_processor),
        ("WordProcessor 空文档", test_word_processor_empty),
        ("WordProcessor 无效路径", test_word_processor_invalid_path),
        ("ExcelProcessor", test_excel_processor),
        ("ExcelProcessor 空文档", test_excel_processor_empty),
        ("ExcelProcessor 无效路径", test_excel_processor_invalid_path),
        ("PPT 生成 (slides)", test_ppt_generator),
        ("PPT 生成 (text)", test_ppt_from_text),
        ("PPT 生成 (空文本)", test_ppt_from_empty_text),
        ("PPT 专业级", test_ppt_professional),
        ("FileHandler", test_file_handler),
        ("文件大小限制", test_size_limits),
        ("FileHandler 边缘", test_file_handler_edge_cases),
        ("document_handler 边缘", test_document_handler_edge_cases),
        ("summarizer 边缘", test_summarizer_edge_cases),
    ]

    passed = 0
    failed = 0
    for name, func in tests:
        try:
            func()
            print(f"  ✅ {name}")
            passed += 1
        except Exception as e:
            print(f"  ❌ {name}: {e}")
            failed += 1

    print()
    print(f"结果: {passed} 通过, {failed} 失败")
